# -*- coding: utf-8 -*-
"""
批量生成复杂测试数据，不低于 100 个文件
格式：png / txt / docx / pdf / xlsx / xls / csv
输出目录：test_data/
"""
import os
import random
from pathlib import Path
from io import BytesIO

BASE_DIR = Path(__file__).parent.parent
OUTPUT_DIR = BASE_DIR / "test_data"
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

random.seed(42)

# ============================================================
# 数据模板库
# ============================================================
COMPANIES = [
    ("星辰科技有限公司", "物联网与边缘计算"),
    ("瀚海数据集团", "大数据与AI分析"),
    ("凌云智造股份", "智能制造与工业4.0"),
    ("瑞丰金融科技", "金融科技与区块链"),
    ("绿源环保科技", "环保监测与新能源"),
]
DEPARTMENTS = ["研发部", "市场部", "销售部", "运营部", "财务部", "人事部", "行政部", "产品部", "技术支持部", "质量管理部"]
PRODUCTS_POOL = [
    ("智能网关", 28500), ("云平台服务", 158000), ("数据分析引擎", 98000),
    ("物联网模组", 4500), ("安全审计系统", 65000), ("边缘计算节点", 32000),
    ("AI推理加速卡", 128000), ("数字孪生平台", 220000), ("工业路由器", 8500),
    ("传感器套件", 3200), ("视频分析系统", 78000), ("区块链存证平台", 55000),
]
CUSTOMERS_POOL = [
    "XX市水务集团", "XX省电力公司", "XX精密制造", "XX智慧城市运营",
    "XX物流科技", "XX医疗设备", "XX环保科技", "XX农业科技",
    "XX交通集团", "XX能源公司", "XX教育科技", "XX零售集团",
]
REGIONS = ["华北", "华东", "华南", "华中", "西南", "西北", "东北"]
REPORT_TOPICS = [
    ("年度经营分析报告", "公司整体经营状况、财务数据、战略回顾"),
    ("市场调研分析报告", "行业趋势、竞品分析、市场份额、用户画像"),
    ("产品技术白皮书", "技术架构、核心算法、性能指标、部署方案"),
    ("项目验收总结报告", "项目背景、实施过程、验收结果、经验总结"),
    ("人力资源分析报告", "人员结构、招聘数据、培训成果、绩效分析"),
    ("财务审计报告", "资产负债表、利润表、现金流量表、审计意见"),
    ("客户满意度调查报告", "调查方法、样本分布、满意度评分、改进建议"),
    ("供应链管理报告", "供应商评估、采购数据、库存分析、物流效率"),
    ("研发创新报告", "专利统计、研发投入、创新成果、技术路线图"),
    ("风险管理评估报告", "风险识别、评估矩阵、应对措施、监控机制"),
    ("数字化转型规划", "现状评估、目标架构、实施路线、投资预算"),
    ("ESG可持续发展报告", "环境指标、社会责任、公司治理、碳中和路径"),
    ("销售业绩分析报告", "区域销售、产品线销售、客户分析、预测模型"),
    ("质量管理体系报告", "质量标准、检测数据、不合格品分析、改进措施"),
    ("知识产权分析报告", "专利布局、商标统计、著作权、技术秘密"),
]
CHART_TYPES = [
    ("月度营收与成本对比", "bar", ["营收（万元）", "成本（万元）"], ["#4CAF50", "#FF9800"]),
    ("部门人员分布", "pie", None, ["#2196F3", "#4CAF50", "#FF9800", "#9C27B0", "#F44336", "#00BCD4", "#795548", "#607D8B", "#E91E63", "#3F51B5"]),
    ("季度利润趋势对比", "line", ["2024年", "2025年"], ["#607D8B", "#E91E63"]),
    ("产品市场份额", "barh", None, ["#2196F3", "#4CAF50", "#FF9800", "#9C27B0", "#F44336", "#00BCD4"]),
    ("客户行业分布", "pie", None, ["#FF5722", "#009688", "#673AB7", "#FFC107", "#00BCD4", "#8BC34A", "#E91E63"]),
    ("月度费用趋势", "line", ["办公费用", "研发费用", "市场费用"], ["#F44336", "#2196F3", "#4CAF50"]),
    ("区域销售对比", "bar", ["2024年", "2025年"], ["#9C27B0", "#FF9800"]),
    ("员工学历分布", "pie", None, ["#3F51B5", "#009688", "#FF9800", "#F44336", "#607D8B"]),
    ("项目完成率", "barh", None, ["#4CAF50", "#FF9800", "#F44336"]),
    ("研发投入趋势", "line", ["研发投入（万元）", "研发占比（%）"], ["#E91E63", "#2196F3"]),
    ("客户满意度评分", "bar", ["2024年", "2025年"], ["#00BCD4", "#FF5722"]),
    ("设备利用率", "barh", None, ["#4CAF50", "#8BC34A", "#CDDC39", "#FFC107", "#FF9800"]),
    ("月度新增客户", "line", ["企业客户", "个人客户"], ["#673AB7", "#FF9800"]),
    ("成本结构分析", "pie", None, ["#F44336", "#2196F3", "#4CAF50", "#FF9800", "#9C27B0"]),
    ("年度KPI达成率", "bar", ["目标值", "实际值"], ["#9E9E9E", "#4CAF50"]),
]


# ============================================================
# 字体初始化
# ============================================================
def _init_matplotlib_font():
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    import matplotlib.font_manager as fm
    font_paths = [
        "C:/Windows/Fonts/simhei.ttf",
        "C:/Windows/Fonts/msyh.ttc",
        "C:/Windows/Fonts/simsun.ttc",
    ]
    for fp in font_paths:
        if os.path.exists(fp):
            fm.fontManager.addfont(fp)
            font_prop = fm.FontProperties(fname=fp)
            font_name = font_prop.get_name()
            plt.rcParams["font.family"] = font_name
            plt.rcParams["axes.unicode_minus"] = False
            return font_name
    return None


def _init_pdf_font():
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    system_fonts = [
        ("C:/Windows/Fonts/simhei.ttf", "SimHei"),
        ("C:/Windows/Fonts/simsun.ttc", "SimSun"),
        ("C:/Windows/Fonts/msyh.ttc", "MSYH"),
    ]
    for fpath, fkey in system_fonts:
        if os.path.exists(fpath):
            try:
                pdfmetrics.registerFont(TTFont(fkey, fpath))
                return fkey, fkey
            except Exception:
                pass
    return "Helvetica", "Helvetica-Bold"


# ============================================================
# 1. 生成 PNG 图表（15 个）
# ============================================================
def generate_png_batch():
    import matplotlib.pyplot as plt
    import numpy as np

    _init_matplotlib_font()
    months = [f"{i}月" for i in range(1, 13)]
    quarters = ["Q1", "Q2", "Q3", "Q4"]

    for idx, (title, ctype, labels, colors) in enumerate(CHART_TYPES):
        fig, ax = plt.subplots(figsize=(8, 5))
        company, _ = COMPANIES[idx % len(COMPANIES)]
        year = 2020 + idx % 6

        if ctype == "bar":
            x = np.arange(len(months))
            width = 0.35
            series1 = [random.randint(80, 300) for _ in range(12)]
            series2 = [random.randint(50, 200) for _ in range(12)]
            ax.bar(x - width / 2, series1, width, label=labels[0], color=colors[0])
            ax.bar(x + width / 2, series2, width, label=labels[1], color=colors[1])
            ax.set_xticks(x)
            ax.set_xticklabels(months, rotation=45)
            ax.legend(fontsize=9)

        elif ctype == "pie":
            parts = random.randint(5, 8)
            values = [random.randint(10, 100) for _ in range(parts)]
            part_labels = DEPARTMENTS[:parts] if "部门" in title or "学历" in title else PRODUCTS_POOL[:parts]
            if isinstance(part_labels[0], tuple):
                part_labels = [p[0] for p in part_labels]
            ax.pie(values, labels=part_labels, colors=colors[:parts],
                   autopct="%1.1f%%", startangle=90)

        elif ctype == "line":
            s1 = [random.randint(150, 400) for _ in range(4)]
            s2 = [random.randint(180, 450) for _ in range(4)]
            ax.plot(quarters, s1, "o-", label=labels[0], color=colors[0], linewidth=2, markersize=8)
            ax.plot(quarters, s2, "s-", label=labels[1], color=colors[1], linewidth=2, markersize=8)
            for i, (v1, v2) in enumerate(zip(s1, s2)):
                ax.annotate(str(v1), (i, v1), textcoords="offset points", xytext=(0, -15), ha="center", fontsize=8)
                ax.annotate(str(v2), (i, v2), textcoords="offset points", xytext=(0, 10), ha="center", fontsize=8)
            ax.legend(fontsize=9)

        elif ctype == "barh":
            parts = random.randint(5, 8)
            items = PRODUCTS_POOL[:parts]
            item_names = [p[0] for p in items]
            values = [random.randint(5, 35) for _ in range(parts)]
            y_pos = np.arange(len(item_names))
            bars = ax.barh(y_pos, values, color=colors[:parts])
            ax.set_yticks(y_pos)
            ax.set_yticklabels(item_names)
            for bar, val in zip(bars, values):
                ax.text(bar.get_width() + 0.3, bar.get_y() + bar.get_height() / 2, str(val), va="center", fontsize=9)

        ax.set_title(f"{company} {year}年{title}")
        ax.grid(alpha=0.3)
        plt.tight_layout()
        buf = BytesIO()
        fig.savefig(buf, format="png", dpi=150)
        plt.close(fig)

        fname = f"chart_{idx+1:02d}_{title}.png"
        with open(OUTPUT_DIR / fname, "wb") as f:
            f.write(buf.getvalue())
        print(f"[PNG {idx+1:02d}/15] {fname}")


# ============================================================
# 2. 生成 TXT 报告（15 个）
# ============================================================
def generate_txt_batch():
    for idx, (topic, desc) in enumerate(REPORT_TOPICS):
        company, industry = COMPANIES[idx % len(COMPANIES)]
        year = 2020 + idx % 6
        revenue = random.randint(800, 5000)
        profit = random.randint(100, 800)
        employees = random.randint(50, 500)
        growth = round(random.uniform(5, 35), 1)

        content = f"""{'='*60}
{company} {year}年{topic}
{'='*60}

编制单位：{company}
报告期间：{year}年1月1日 - {year}年12月31日
行业领域：{industry}
密级：内部机密

第一章  概述

1.1 报告背景
本报告由{company}编制，旨在全面反映公司在{year}年度的{desc}。
公司成立于2018年，总部位于深圳市南山区科技园，是一家专注于{industry}
领域的高新技术企业。截至{year}年末，公司拥有员工{employees}人。

1.2 核心经营指标
  (1) 年度营收：{revenue}万元（同比增长{growth}%）
  (2) 净利润：{profit}万元（净利率{round(profit/revenue*100, 1)}%）
  (3) 员工总数：{employees}人
  (4) 研发投入：{random.randint(50, 500)}万元（占比{round(random.uniform(8, 25), 1)}%）
  (5) 新增客户：{random.randint(20, 150)}家
  (6) 专利申请：{random.randint(5, 30)}项

第二章  详细分析

2.1 市场环境分析
{year}年，{industry}行业持续快速发展。根据工信部数据，行业整体规模
达到{random.randint(500, 5000)}亿元，同比增长{round(random.uniform(10, 30), 1)}%。
公司抓住行业数字化转型机遇，在以下方面取得突破：
  - 产品技术迭代：完成{random.randint(2, 6)}个重大版本升级
  - 市场拓展：新进入{random.randint(1, 4)}个细分领域
  - 生态合作：与{random.randint(3, 10)}家头部企业建立战略合作

2.2 财务数据分析
                                                     单位：万元
+------------------+----------+----------+----------+----------+
|     指标         |   Q1     |   Q2     |   Q3     |   Q4     |
+------------------+----------+----------+----------+----------+
| 营业收入         |{random.randint(150,800):>8}  |{random.randint(180,900):>8}  |{random.randint(200,1000):>8}  |{random.randint(250,1200):>8}  |
| 营业成本         |{random.randint(80,500):>8}  |{random.randint(90,550):>8}  |{random.randint(100,600):>8}  |{random.randint(120,700):>8}  |
| 毛利润           |{random.randint(50,300):>8}  |{random.randint(60,350):>8}  |{random.randint(70,400):>8}  |{random.randint(80,500):>8}  |
| 研发费用         |{random.randint(20,150):>8}  |{random.randint(25,160):>8}  |{random.randint(30,180):>8}  |{random.randint(35,200):>8}  |
| 净利润           |{random.randint(10,100):>8}  |{random.randint(15,120):>8}  |{random.randint(20,150):>8}  |{random.randint(25,180):>8}  |
+------------------+----------+----------+----------+----------+

2.3 SWOT分析
  【优势 Strengths】
    - 核心技术自主可控，拥有{random.randint(10, 50)}项发明专利
    - 团队经验丰富，核心成员平均从业{random.randint(8, 15)}年
    - 客户粘性高，续约率达{round(random.uniform(75, 95), 1)}%

  【劣势 Weaknesses】
    - 品牌知名度在部分区域市场有待提升
    - 国际化布局尚处于起步阶段
    - 高端人才招聘竞争激烈

  【机会 Opportunities】
    - 国家政策大力支持{industry}发展
    - 5G/AI/云计算技术融合带来新场景
    - 海外新兴市场需求快速增长

  【威胁 Threats】
    - 行业竞争加剧，价格战压力增大
    - 技术迭代速度快，研发投入压力大
    - 宏观经济不确定性增加

第三章  数据附表

3.1 产品线营收明细
"""
        for pi, (pname, price) in enumerate(PRODUCTS_POOL[:8]):
            qty = random.randint(10, 500)
            rev = qty * price // 10000
            content += f"  {pi+1}. {pname}：销售{qty}套，营收{rev}万元\n"

        content += f"""
3.2 区域销售分布
"""
        for region in REGIONS:
            content += f"  {region}地区：{random.randint(50, 600)}万元（占比{round(random.uniform(5, 25), 1)}%）\n"

        content += f"""
第四章  总结与展望

4.1 年度总结
{year}年是公司发展的关键一年。在全体员工的共同努力下，公司克服了
各种挑战，取得了令人瞩目的成绩。营收首次突破{revenue}万元大关，
同比增长{growth}%，远超行业平均水平。

4.2 下一年度规划
  (1) 营收目标：{int(revenue * (1 + random.uniform(0.2, 0.5)))}万元
  (2) 团队扩充至{int(employees * (1 + random.uniform(0.1, 0.3)))}人
  (3) 启动C轮融资，目标{random.randint(5000, 20000)}万元
  (4) 拓展海外市场，设立{random.randint(1, 3)}个海外办事处
  (5) 加大AI研发投入，建设联合实验室

---
报告编制人：{company}经营管理部
编制日期：{year}年12月31日
联系方式：report@{company.replace('有限公司', '').replace('股份', '').replace('集团', '').lower()}.com
官网：https://www.example-{idx+1:02d}.com
"""

        fname = f"report_{idx+1:02d}_{topic}.txt"
        with open(OUTPUT_DIR / fname, "w", encoding="utf-8") as f:
            f.write(content)
        print(f"[TXT {idx+1:02d}/15] {fname} ({len(content)} 字符)")


# ============================================================
# 3. 生成 DOCX 文档（15 个）
# ============================================================
def generate_docx_batch():
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT

    for idx in range(15):
        company, industry = COMPANIES[idx % len(COMPANIES)]
        year = 2020 + idx % 6
        doc = Document()

        style = doc.styles["Normal"]
        style.font.name = "微软雅黑"
        style.font.size = Pt(10.5)

        # 封面
        for _ in range(6):
            doc.add_paragraph()
        title_p = doc.add_paragraph()
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title_p.add_run(f"{company}\n{year}年度产品技术手册")
        run.font.size = Pt(26)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x1A, 0x73, 0xE8)

        doc.add_paragraph()
        sub_p = doc.add_paragraph()
        sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub_p.add_run(f"行业领域：{industry}").font.size = Pt(14)

        doc.add_page_break()

        # 目录页
        doc.add_heading("目录", level=1)
        for i, section in enumerate(["产品概述", "技术架构", "核心功能", "性能指标",
                                       "部署方案", "安全体系", "API接口", "运维管理",
                                       "案例分析", "版本历史"]):
            doc.add_paragraph(f"{i+1}. {section}")

        doc.add_page_break()

        # 第1章
        doc.add_heading(f"1. 产品概述", level=1)
        doc.add_paragraph(
            f"本手册详细介绍了{company}在{industry}领域的核心产品体系。"
            f"公司自成立以来，始终坚持以技术创新为驱动，为客户提供高品质的解决方案。"
            f"截至{year}年，公司已服务超过{random.randint(100, 1000)}家企业客户，"
            f"产品覆盖{', '.join(random.sample(REGIONS, 3))}等主要区域市场。"
        )

        # 第2章
        doc.add_heading("2. 技术架构", level=1)
        doc.add_paragraph(
            "系统采用微服务架构设计，基于Kubernetes容器编排平台，支持弹性伸缩和高可用部署。"
            "核心技术栈包括：Spring Cloud微服务框架、Apache Kafka消息队列、"
            "Redis缓存集群、PostgreSQL关系数据库、Elasticsearch搜索引擎。"
        )

        doc.add_heading("2.1 架构分层", level=2)
        layers = [
            ("接入层", "Nginx反向代理、API Gateway、负载均衡、WAF防火墙"),
            ("应用层", "业务微服务集群、工作流引擎、规则引擎、消息中心"),
            ("数据层", "分布式数据库、对象存储、数据仓库、实时计算引擎"),
            ("基础设施层", "Kubernetes集群、服务网格、监控告警、日志收集"),
        ]
        for lname, ldesc in layers:
            doc.add_paragraph(f"{lname}：{ldesc}", style="List Bullet")

        # 第3章 - 含表格
        doc.add_heading("3. 核心功能", level=1)
        doc.add_paragraph("以下是产品核心功能模块的详细说明：")

        table = doc.add_table(rows=9, cols=4)
        table.style = "Light Grid Accent 1"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        headers = ["功能模块", "功能描述", "适用场景", "版本要求"]
        for ci, h in enumerate(headers):
            table.rows[0].cells[ci].text = h

        features = [
            ("实时数据采集", "支持Modbus/OPC-UA/MQTT等协议", "工业物联网", "V3.0+"),
            ("智能告警引擎", "基于规则+ML的多级告警", "运维监控", "V2.5+"),
            ("可视化大屏", "拖拽式图表配置，实时刷新", "指挥中心", "V3.2+"),
            ("数据分析平台", "SQL查询+Python脚本分析", "数据中台", "V3.0+"),
            ("设备管理", "全生命周期管理，远程配置", "资产管理", "V2.0+"),
            ("权限管理", "RBAC+ABAC混合权限模型", "安全管理", "V1.0+"),
            ("报表中心", "自定义报表模板，定时推送", "业务分析", "V2.8+"),
            ("开放API", "RESTful/GraphQL/gRPC接口", "系统集成", "V3.0+"),
        ]
        for ri, (mod, desc, scene, ver) in enumerate(features):
            row = table.rows[ri + 1]
            row.cells[0].text = mod
            row.cells[1].text = desc
            row.cells[2].text = scene
            row.cells[3].text = ver

        # 第4章
        doc.add_heading("4. 性能指标", level=1)
        perf_data = [
            ("并发连接数", f"{random.randint(10000, 100000)}+"),
            ("数据吞吐量", f"{random.randint(100, 1000)}MB/s"),
            ("平均响应时间", f"< {random.randint(10, 100)}ms"),
            ("系统可用性", "99.99%"),
            ("数据准确率", f"99.{random.randint(9, 99)}%"),
            ("单节点设备接入", f"{random.randint(1000, 10000)}+"),
        ]
        for metric, value in perf_data:
            doc.add_paragraph(f"{metric}：{value}", style="List Bullet")

        # 第5章
        doc.add_heading("5. 部署方案", level=1)
        doc.add_paragraph("支持以下三种部署模式：")
        deploy_modes = [
            ("私有化部署", "适用于对数据安全要求高的企业，部署在客户自有数据中心"),
            ("云托管部署", "部署在阿里云/华为云/腾讯云，由我方负责运维"),
            ("混合云部署", "核心数据本地存储，计算资源弹性使用公有云"),
        ]
        for mode, desc in deploy_modes:
            doc.add_heading(mode, level=2)
            doc.add_paragraph(desc)

        # 第6章
        doc.add_heading("6. 安全体系", level=1)
        doc.add_paragraph(
            "系统通过ISO 27001信息安全管理体系认证，具备完善的安全防护能力：\n"
            "  - 传输安全：全链路TLS 1.3加密\n"
            "  - 存储安全：AES-256数据加密\n"
            "  - 访问控制：多因素认证 + 细粒度权限\n"
            "  - 审计日志：全操作记录，不可篡改\n"
            "  - 漏洞管理：定期渗透测试 + 安全扫描"
        )

        # 第7章
        doc.add_heading("7. API接口", level=1)
        doc.add_paragraph("系统提供丰富的RESTful API接口，支持JSON/Protobuf格式：")
        apis = [
            ("GET", "/api/v1/devices", "获取设备列表"),
            ("POST", "/api/v1/devices", "注册新设备"),
            ("GET", "/api/v1/data/query", "查询历史数据"),
            ("POST", "/api/v1/alarm/subscribe", "订阅告警通知"),
            ("GET", "/api/v1/report/generate", "生成分析报告"),
        ]
        for method, path, desc in apis:
            doc.add_paragraph(f"{method} {path} - {desc}", style="List Bullet")

        # 第8章
        doc.add_heading("8. 运维管理", level=1)
        doc.add_paragraph(
            "提供全方位的运维管理工具：\n"
            "  - 监控面板：Grafana + Prometheus 实时监控\n"
            "  - 日志管理：ELK Stack 集中日志分析\n"
            "  - 备份恢复：自动备份策略，支持异地容灾\n"
            "  - 版本升级：灰度发布，支持热更新和回滚"
        )

        # 第9章
        doc.add_heading("9. 案例分析", level=1)
        for ci in range(3):
            cust = CUSTOMERS_POOL[ci]
            doc.add_heading(f"9.{ci+1} {cust}", level=2)
            doc.add_paragraph(
                f"{cust}于{year-1}年引入我司产品，部署{random.randint(10, 200)}个"
                f"监测节点，覆盖{random.randint(3, 20)}个业务场景。上线后：\n"
                f"  - 运维效率提升{random.randint(30, 80)}%\n"
                f"  - 故障响应时间缩短至{random.randint(1, 15)}分钟\n"
                f"  - 年度节省成本约{random.randint(50, 500)}万元"
            )

        # 第10章
        doc.add_heading("10. 版本历史", level=1)
        for vi in range(5):
            ver_year = year - 4 + vi
            doc.add_paragraph(
                f"V{vi+1}.0（{ver_year}年）：{random.choice(['重大架构升级', '新增AI分析引擎', '性能大幅优化', '新增移动端支持', '国际化多语言支持'])}"
            )

        # 联系方式
        doc.add_page_break()
        doc.add_heading("联系我们", level=1)
        doc.add_paragraph(f"公司名称：{company}")
        doc.add_paragraph(f"地址：深圳市南山区科技园南区高新中三道{random.randint(1, 20)}号")
        doc.add_paragraph(f"电话：0755-{random.randint(1000, 9999)}-{random.randint(1000, 9999)}")
        doc.add_paragraph(f"邮箱：info@{company.replace('有限公司', '').replace('股份', '').replace('集团', '').lower()}.com")
        doc.add_paragraph(f"官网：https://www.example-{idx+1:02d}.com")

        fname = f"doc_{idx+1:02d}_产品技术手册_{company}.docx"
        doc.save(str(OUTPUT_DIR / fname))
        print(f"[DOCX {idx+1:02d}/15] {fname}")


# ============================================================
# 4. 生成 PDF 报告（15 个）
# ============================================================
def generate_pdf_batch():
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import mm, cm
    from reportlab.lib.colors import HexColor, white, grey
    from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak, HRFlowable
    )

    font_name, font_bold = _init_pdf_font()

    for idx in range(15):
        company, industry = COMPANIES[idx % len(COMPANIES)]
        year = 2020 + idx % 6
        topic, desc = REPORT_TOPICS[idx]

        fname = f"pdf_{idx+1:02d}_{topic}_{company}.pdf"
        path = OUTPUT_DIR / fname
        doc = SimpleDocTemplate(str(path), pagesize=A4,
                                leftMargin=2*cm, rightMargin=2*cm,
                                topMargin=2*cm, bottomMargin=2*cm)
        styles = getSampleStyleSheet()

        def ms(name, parent, **kw):
            return ParagraphStyle(name, parent=styles[parent], **kw)

        title_s = ms("T", "Title", fontName=font_bold, fontSize=22,
                     spaceAfter=12, alignment=TA_CENTER)
        h1_s = ms("H1", "Heading1", fontName=font_bold, fontSize=16,
                  spaceBefore=16, spaceAfter=8)
        h2_s = ms("H2", "Heading2", fontName=font_bold, fontSize=13,
                  spaceBefore=12, spaceAfter=6)
        body_s = ms("B", "Normal", fontName=font_name, fontSize=10,
                    leading=16, spaceAfter=6, alignment=TA_JUSTIFY)
        tc_s = ParagraphStyle("TC", fontName=font_name, fontSize=8,
                              leading=12, alignment=TA_CENTER)
        th_s = ParagraphStyle("TH", fontName=font_bold, fontSize=8,
                              leading=12, alignment=TA_CENTER)

        def mt(data, col_widths, header_bg=HexColor("#1A73E8")):
            rows = [[Paragraph(c, th_s if i == 0 else tc_s)
                     for i, c in enumerate(r)] for r in data]
            t = Table(rows, colWidths=col_widths)
            t.setStyle(TableStyle([
                ("GRID", (0, 0), (-1, -1), 0.5, grey),
                ("BACKGROUND", (0, 0), (-1, 0), header_bg),
                ("TEXTCOLOR", (0, 0), (-1, 0), white),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("TOPPADDING", (0, 0), (-1, -1), 3),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
            ]))
            return t

        story = []

        # 封面
        story.append(Spacer(1, 60*mm))
        story.append(Paragraph(company, title_s))
        story.append(Spacer(1, 10*mm))
        story.append(Paragraph(f"{year}年{topic}", ms("Sub", "Title",
                            fontName=font_bold, fontSize=14, alignment=TA_CENTER)))
        story.append(Spacer(1, 20*mm))
        story.append(HRFlowable(width="60%", thickness=1, color=HexColor("#1A73E8")))
        story.append(Spacer(1, 10*mm))
        for line in [f"行业领域：{industry}", f"报告期间：{year}年度",
                     f"编制单位：{company}", f"密级：内部机密"]:
            story.append(Paragraph(line, ms("I", "Normal", fontName=font_name,
                                            fontSize=11, alignment=TA_CENTER)))
        story.append(PageBreak())

        # 第1章
        story.append(Paragraph("1. 报告概述", h1_s))
        story.append(Paragraph(
            f"本报告由{company}编制，全面反映公司在{year}年度的{desc}。"
            f"报告数据来源于公司ERP系统、财务系统和各业务部门统计报表，"
            f"经内部审计部门审核确认。", body_s))

        story.append(Paragraph("1.1 基本信息", h2_s))
        story.append(mt([
            ["项目", "内容"],
            ["编制单位", company],
            ["行业领域", industry],
            ["报告年度", f"{year}年"],
            ["员工总数", f"{random.randint(50, 500)}人"],
            ["年度营收", f"{random.randint(800, 5000)}万元"],
        ], [4*cm, 10*cm]))
        story.append(Spacer(1, 6*mm))

        # 第2章
        story.append(Paragraph("2. 经营数据分析", h1_s))
        story.append(Paragraph("2.1 季度营收明细", h2_s))
        qdata = [["季度", "营收（万元）", "成本（万元）", "利润（万元）", "利润率"]]
        for qi, qname in enumerate(["Q1", "Q2", "Q3", "Q4"]):
            rev = random.randint(150, 1200)
            cost = random.randint(80, 700)
            prof = rev - cost
            qdata.append([qname, str(rev), str(cost), str(prof), f"{round(prof/rev*100, 1)}%"])
        story.append(mt(qdata, [2.5*cm, 3*cm, 3*cm, 3*cm, 2.5*cm]))
        story.append(Spacer(1, 6*mm))

        # 第3章
        story.append(Paragraph("3. 产品线分析", h1_s))
        pdata = [["产品名称", "销售额（万元）", "同比增长", "市场份额"]]
        for pi in range(6):
            pname = PRODUCTS_POOL[pi][0]
            pdata.append([
                pname,
                str(random.randint(50, 800)),
                f"{round(random.uniform(-5, 40), 1)}%",
                f"{round(random.uniform(5, 30), 1)}%",
            ])
        story.append(mt(pdata, [3.5*cm, 3*cm, 3*cm, 3*cm]))
        story.append(Spacer(1, 6*mm))

        # 第4章
        story.append(Paragraph("4. 风险评估与应对", h1_s))
        risks = [
            ("市场风险", "行业竞争加剧，价格战可能导致利润率下降",
             "加强差异化竞争，提升产品附加值"),
            ("技术风险", "技术迭代快，现有产品可能被替代",
             "加大研发投入，保持技术领先"),
            ("人才风险", "高端人才招聘困难，核心人员流失",
             "完善激励机制，建立人才梯队"),
            ("政策风险", "行业监管政策变化可能影响业务",
             "密切关注政策动向，提前布局合规"),
        ]
        rdata = [["风险类别", "风险描述", "应对措施"]]
        for rc, rd, rm in risks:
            rdata.append([rc, rd, rm])
        story.append(mt(rdata, [2.5*cm, 5*cm, 5*cm]))
        story.append(Spacer(1, 6*mm))

        # 第5章
        story.append(Paragraph("5. 总结与展望", h1_s))
        story.append(Paragraph(
            f"{year}年，{company}在{industry}领域取得了显著成绩。展望未来，"
            f"公司将继续坚持技术创新驱动，深化行业应用，拓展海外市场，"
            f"力争在{year+1}年实现营收突破{random.randint(1000, 8000)}万元。", body_s))

        story.append(Spacer(1, 10*mm))
        story.append(HRFlowable(width="100%", thickness=0.5, color=grey))
        story.append(Paragraph(f"报告编制：{company} | 日期：{year}年12月31日",
                               ms("F", "Normal", fontName=font_name, fontSize=8, alignment=TA_CENTER)))

        doc.build(story)
        print(f"[PDF {idx+1:02d}/15] {fname}")


# ============================================================
# 5. 生成 XLSX 文件（15 个）
# ============================================================
def generate_xlsx_batch():
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.chart import BarChart, Reference
    from openpyxl.utils import get_column_letter

    for idx in range(15):
        company, industry = COMPANIES[idx % len(COMPANIES)]
        year = 2020 + idx % 6
        wb = Workbook()

        # 样式
        tfont = Font(name="微软雅黑", size=14, bold=True, color="1A73E8")
        hfont = Font(name="微软雅黑", size=10, bold=True, color="FFFFFF")
        hfill = PatternFill(start_color="1A73E8", end_color="1A73E8", fill_type="solid")
        dfont = Font(name="微软雅黑", size=10)
        thin_border = Border(
            left=Side(style="thin"), right=Side(style="thin"),
            top=Side(style="thin"), bottom=Side(style="thin"),
        )
        center_align = Alignment(horizontal="center", vertical="center")

        def style_header(ws, row, cols):
            for c in range(1, cols + 1):
                cell = ws.cell(row=row, column=c)
                cell.font = hfont
                cell.fill = hfill
                cell.alignment = center_align
                cell.border = thin_border

        def style_data(ws, row, cols):
            for c in range(1, cols + 1):
                cell = ws.cell(row=row, column=c)
                cell.font = dfont
                cell.border = thin_border
                cell.alignment = center_align

        # Sheet 1: 财务总表
        ws1 = wb.active
        ws1.title = "财务总表"
        ws1.merge_cells("A1:G1")
        ws1["A1"] = f"{company} {year}年度财务总表"
        ws1["A1"].font = tfont
        ws1["A1"].alignment = Alignment(horizontal="center", vertical="center")

        headers1 = ["月份", "营收（万元）", "成本（万元）", "毛利润（万元）",
                     "费用（万元）", "净利润（万元）", "净利率"]
        for c, h in enumerate(headers1, 1):
            ws1.cell(row=3, column=c, value=h)
        style_header(ws1, 3, 7)

        for m in range(1, 13):
            r = m + 3
            rev = random.randint(150, 350)
            cost = random.randint(80, 200)
            fee = random.randint(20, 80)
            ws1.cell(row=r, column=1, value=f"{m}月")
            ws1.cell(row=r, column=2, value=rev)
            ws1.cell(row=r, column=3, value=cost)
            ws1.cell(row=r, column=4).value = f"=B{r}-C{r}"
            ws1.cell(row=r, column=5, value=fee)
            ws1.cell(row=r, column=6).value = f"=D{r}-E{r}"
            ws1.cell(row=r, column=7).value = f'=IF(B{r}=0,"",F{r}/B{r})'
            ws1.cell(row=r, column=7).number_format = "0.0%"
            style_data(ws1, r, 7)

        # 合计行
        sr = 16
        ws1.cell(row=sr, column=1, value="合计")
        ws1.cell(row=sr, column=1).font = Font(name="微软雅黑", size=10, bold=True)
        for c in range(2, 7):
            col = get_column_letter(c)
            ws1.cell(row=sr, column=c).value = f"=SUM({col}4:{col}15)"
            ws1.cell(row=sr, column=c).font = Font(name="微软雅黑", size=10, bold=True)

        # 图表
        chart = BarChart()
        chart.title = f"{company} 月度营收与利润"
        chart.y_axis.title = "金额（万元）"
        chart.x_axis.title = "月份"
        data_ref = Reference(ws1, min_col=2, max_col=6, min_row=3, max_row=15)
        cats_ref = Reference(ws1, min_col=1, min_row=4, max_row=15)
        chart.add_data(data_ref, titles_from_data=True)
        chart.set_categories(cats_ref)
        chart.width = 20
        chart.height = 12
        ws1.add_chart(chart, "A18")

        for c, w in enumerate([8, 14, 14, 14, 14, 14, 10], 1):
            ws1.column_dimensions[get_column_letter(c)].width = w

        # Sheet 2: 销售明细
        ws2 = wb.create_sheet("销售明细")
        ws2.merge_cells("A1:G1")
        ws2["A1"] = f"{company} 销售明细表"
        ws2["A1"].font = tfont
        ws2["A1"].alignment = center_align

        headers2 = ["日期", "产品名称", "客户名称", "数量", "单价", "金额", "区域"]
        for c, h in enumerate(headers2, 1):
            ws2.cell(row=3, column=c, value=h)
        style_header(ws2, 3, 7)

        for r in range(4, 54):
            m = random.randint(1, 12)
            d = random.randint(1, 28)
            pi = random.randint(0, len(PRODUCTS_POOL) - 1)
            pname, price = PRODUCTS_POOL[pi]
            qty = random.randint(1, 50)
            ws2.cell(row=r, column=1, value=f"{year}-{m:02d}-{d:02d}")
            ws2.cell(row=r, column=2, value=pname)
            ws2.cell(row=r, column=3, value=random.choice(CUSTOMERS_POOL))
            ws2.cell(row=r, column=4, value=qty)
            ws2.cell(row=r, column=5, value=price)
            ws2.cell(row=r, column=5).number_format = "#,##0"
            ws2.cell(row=r, column=6).value = f"=D{r}*E{r}"
            ws2.cell(row=r, column=6).number_format = "#,##0"
            ws2.cell(row=r, column=7, value=random.choice(REGIONS))
            style_data(ws2, r, 7)

        for c, w in enumerate([14, 16, 18, 8, 10, 12, 8], 1):
            ws2.column_dimensions[get_column_letter(c)].width = w

        # Sheet 3: 员工信息
        ws3 = wb.create_sheet("员工信息")
        ws3.merge_cells("A1:H1")
        ws3["A1"] = f"{company} 员工信息表"
        ws3["A1"].font = tfont
        ws3["A1"].alignment = center_align

        headers3 = ["工号", "姓名", "部门", "职位", "入职日期", "学历", "薪资", "绩效"]
        for c, h in enumerate(headers3, 1):
            ws3.cell(row=3, column=c, value=h)
        style_header(ws3, 3, 8)

        surnames = "张李王刘陈杨赵黄周吴徐孙马胡朱郭何罗高林郑梁谢唐许冯宋韩邓彭曹"
        dept_pool = DEPARTMENTS[:7]
        positions = ["工程师", "高级工程师", "经理", "主管", "专员", "架构师", "总监"]
        degrees = ["博士", "硕士", "本科", "大专"]

        for r in range(4, 34):
            ws3.cell(row=r, column=1, value=f"EMP{idx*100+r:04d}")
            ws3.cell(row=r, column=2, value=random.choice(surnames) + random.choice("明华伟芳丽强军敏静"))
            ws3.cell(row=r, column=3, value=random.choice(dept_pool))
            ws3.cell(row=r, column=4, value=random.choice(positions))
            ws3.cell(row=r, column=5, value=f"{year-random.randint(1,8)}-{random.randint(1,12):02d}-{random.randint(1,28):02d}")
            ws3.cell(row=r, column=6, value=random.choice(degrees))
            ws3.cell(row=r, column=7, value=random.randint(8000, 50000))
            ws3.cell(row=r, column=7).number_format = "#,##0"
            ws3.cell(row=r, column=8, value=random.choice(["A", "B+", "B", "C", "A+"]))
            style_data(ws3, r, 8)

        for c, w in enumerate([10, 10, 10, 12, 14, 8, 10, 8], 1):
            ws3.column_dimensions[get_column_letter(c)].width = w

        fname = f"xlsx_{idx+1:02d}_数据报表_{company}.xlsx"
        wb.save(str(OUTPUT_DIR / fname))
        print(f"[XLSX {idx+1:02d}/15] {fname}")


# ============================================================
# 6. 生成 XLS 文件（15 个）
# ============================================================
def generate_xls_batch():
    import xlwt

    for idx in range(15):
        company, industry = COMPANIES[idx % len(COMPANIES)]
        year = 2020 + idx % 6
        wb = xlwt.Workbook(encoding="utf-8")

        # 样式
        title_style = xlwt.XFStyle()
        tf = xlwt.Font()
        tf.name = "微软雅黑"
        tf.height = 320
        tf.bold = True
        tf.colour_index = 0x1E
        title_style.font = tf
        ta = xlwt.Alignment()
        ta.horz = xlwt.Alignment.HORZ_CENTER
        ta.vert = xlwt.Alignment.VERT_CENTER
        title_style.alignment = ta

        header_style = xlwt.XFStyle()
        hf = xlwt.Font()
        hf.name = "微软雅黑"
        hf.height = 220
        hf.bold = True
        hf.colour_index = 1
        header_style.font = hf
        hp = xlwt.Pattern()
        hp.pattern = xlwt.Pattern.SOLID_PATTERN
        hp.pattern_fore_colour = 0x1E
        header_style.pattern = hp
        ha = xlwt.Alignment()
        ha.horz = xlwt.Alignment.HORZ_CENTER
        ha.vert = xlwt.Alignment.VERT_CENTER
        header_style.alignment = ha

        data_style = xlwt.XFStyle()
        df = xlwt.Font()
        df.name = "微软雅黑"
        df.height = 200
        data_style.font = df
        da = xlwt.Alignment()
        da.horz = xlwt.Alignment.HORZ_CENTER
        da.vert = xlwt.Alignment.VERT_CENTER
        data_style.alignment = da

        money_style = xlwt.XFStyle()
        money_style.font = df
        money_style.alignment = da
        money_style.num_format_str = "#,##0.00"

        borders = xlwt.Borders()
        borders.left = xlwt.Borders.THIN
        borders.right = xlwt.Borders.THIN
        borders.top = xlwt.Borders.THIN
        borders.bottom = xlwt.Borders.THIN
        data_style.borders = borders
        money_style.borders = borders
        header_style.borders = borders

        # Sheet 1: 项目成本核算
        ws1 = wb.add_sheet("项目成本核算")
        ws1.write_merge(0, 0, 0, 6, f"{company} {year}年项目成本核算表", title_style)

        headers = ["项目编号", "项目名称", "客户名称", "合同金额", "硬件成本", "人力成本", "毛利润"]
        for c, h in enumerate(headers):
            ws1.write(2, c, h, header_style)

        projects = [
            (f"PRJ-{year}-001", "智慧水务改造一期", "XX市水务集团", 2200000.00, 850000.00, 620000.00),
            (f"PRJ-{year}-002", "智能配电网网关采购", "XX省电力公司", 3800000.00, 1520000.00, 980000.00),
            (f"PRJ-{year}-003", "工业物联网平台部署", "XX精密制造", 1800000.00, 720000.00, 450000.00),
            (f"PRJ-{year}-004", "智慧城市数据中台", "XX智慧城市运营", 1500000.00, 580000.00, 420000.00),
            (f"PRJ-{year}-005", "物流IoT追踪系统", "XX物流科技", 1200000.00, 460000.00, 350000.00),
            (f"PRJ-{year}-006", "医疗设备监控平台", "XX医疗设备", 950000.00, 380000.00, 280000.00),
            (f"PRJ-{year}-007", "环保监测数据平台", "XX环保科技", 800000.00, 320000.00, 240000.00),
            (f"PRJ-{year}-008", "农业物联网系统", "XX农业科技", 650000.00, 260000.00, 190000.00),
            (f"PRJ-{year}-009", "智慧交通边缘计算", "XX交通集团", 2800000.00, 1120000.00, 780000.00),
            (f"PRJ-{year}-010", "能源管理云平台", "XX能源公司", 2100000.00, 840000.00, 560000.00),
        ]

        for i, (pid, pname, cust, contract, hw, labor) in enumerate(projects):
            r = 3 + i
            ws1.write(r, 0, pid, data_style)
            ws1.write(r, 1, pname, data_style)
            ws1.write(r, 2, cust, data_style)
            ws1.write(r, 3, contract, money_style)
            ws1.write(r, 4, hw, money_style)
            ws1.write(r, 5, labor, money_style)
            ws1.write(r, 6, xlwt.Formula(f"D{r+1}-E{r+1}-F{r+1}"), money_style)

        sum_row = 13
        sum_style = xlwt.XFStyle()
        sf = xlwt.Font()
        sf.name = "微软雅黑"
        sf.height = 200
        sf.bold = True
        sum_style.font = sf
        sum_style.borders = borders
        sum_style.num_format_str = "#,##0.00"
        sum_style.alignment = da

        ws1.write(sum_row, 0, "合计", sum_style)
        for c in range(3, 7):
            col_letter = chr(ord("D") + c - 3)
            ws1.write(sum_row, c, xlwt.Formula(f"SUM({col_letter}4:{col_letter}13)"), sum_style)

        for c, w in enumerate([4000, 8000, 6000, 5000, 5000, 5000, 5000]):
            ws1.col(c).width = w

        # Sheet 2: 设备清单
        ws2 = wb.add_sheet("设备清单")
        ws2.write_merge(0, 0, 0, 5, f"{company} 设备资产清单", title_style)

        headers2 = ["资产编号", "设备名称", "型号", "部门", "购置日期", "原值"]
        for c, h in enumerate(headers2):
            ws2.write(2, c, h, header_style)

        devices = [
            ("AST-001", "服务器", "Dell R750", "研发部", "2024-03-15", 85000.00),
            ("AST-002", "服务器", "Dell R750", "研发部", "2024-03-15", 85000.00),
            ("AST-003", "GPU服务器", "NVIDIA DGX A100", "研发部", "2024-06-20", 350000.00),
            ("AST-004", "交换机", "Huawei S6730", "研发部", "2024-01-10", 28000.00),
            ("AST-005", "防火墙", "H3C SecPath F5000", "研发部", "2024-01-10", 45000.00),
            ("AST-006", "笔记本电脑", "ThinkPad X1", "市场部", "2025-02-15", 12000.00),
            ("AST-007", "笔记本电脑", "ThinkPad X1", "市场部", "2025-02-15", 12000.00),
            ("AST-008", "笔记本电脑", "MacBook Pro", "销售部", "2025-03-01", 18000.00),
            ("AST-009", "打印机", "HP LaserJet", "行政部", "2023-08-10", 8500.00),
            ("AST-010", "投影仪", "Epson CB-L730U", "行政部", "2024-05-20", 22000.00),
            ("AST-011", "示波器", "Tektronix MDO34", "研发部", "2024-09-15", 68000.00),
            ("AST-012", "频谱分析仪", "Keysight N9320B", "研发部", "2024-09-15", 95000.00),
            ("AST-013", "3D打印机", "Ultimaker S5", "研发部", "2025-01-10", 42000.00),
            ("AST-014", "空调", "大金商用中央空调", "行政部", "2023-06-01", 120000.00),
            ("AST-015", "UPS电源", "APC Smart-UPS 3000", "研发部", "2024-01-10", 25000.00),
        ]

        for i, (aid, name, model, dept, date, value) in enumerate(devices):
            r = 3 + i
            ws2.write(r, 0, aid, data_style)
            ws2.write(r, 1, name, data_style)
            ws2.write(r, 2, model, data_style)
            ws2.write(r, 3, dept, data_style)
            ws2.write(r, 4, date, data_style)
            ws2.write(r, 5, value, money_style)

        for c, w in enumerate([4000, 6000, 8000, 4000, 5000, 5000]):
            ws2.col(c).width = w

        fname = f"xls_{idx+1:02d}_项目成本_{company}.xls"
        wb.save(str(OUTPUT_DIR / fname))
        print(f"[XLS {idx+1:02d}/15] {fname}")


# ============================================================
# 7. 生成 CSV 文件（10 个）
# ============================================================
def generate_csv_batch():
    import csv

    csv_topics = [
        ("销售订单数据", ["订单号", "日期", "产品", "客户", "数量", "单价", "金额", "区域", "销售员"]),
        ("设备监控数据", ["设备ID", "时间戳", "温度", "湿度", "电压", "电流", "功率", "状态", "告警级别"]),
        ("客户反馈数据", ["反馈ID", "日期", "客户", "产品", "评分", "反馈类型", "处理状态", "处理人", "备注"]),
        ("库存管理数据", ["SKU", "产品名称", "仓库", "库存量", "安全库存", "在途量", "供应商", "上次盘点", "状态"]),
        ("项目进度数据", ["项目ID", "项目名称", "阶段", "完成度", "负责人", "开始日期", "截止日期", "状态", "风险等级"]),
        ("财务流水数据", ["流水号", "日期", "类型", "金额", "科目", "对方单位", "摘要", "经办人", "审批状态"]),
        ("物流运输数据", ["运单号", "发货日期", "目的地", "货物类型", "重量", "体积", "承运商", "运费", "签收状态"]),
        ("质量检测数据", ["检测编号", "日期", "产品批次", "检测项目", "标准值", "实测值", "偏差", "判定", "检测员"]),
        ("培训记录数据", ["培训编号", "日期", "课程名称", "讲师", "参训人数", "通过率", "满意度", "费用", "备注"]),
        ("市场活动数据", ["活动ID", "日期", "活动类型", "主题", "预算", "实际花费", "参与人数", "转化率", "ROI"]),
    ]

    for idx, (topic, headers) in enumerate(csv_topics):
        fname = f"csv_{idx+1:02d}_{topic}.csv"
        with open(OUTPUT_DIR / fname, "w", encoding="utf-8-sig", newline="") as f:
            writer = csv.writer(f)
            writer.writerow(headers)

            for row_idx in range(50):
                if "销售" in topic:
                    pi = random.randint(0, len(PRODUCTS_POOL) - 1)
                    pname, price = PRODUCTS_POOL[pi]
                    qty = random.randint(1, 50)
                    writer.writerow([
                        f"ORD-{idx*1000+row_idx:06d}",
                        f"2025-{random.randint(1,12):02d}-{random.randint(1,28):02d}",
                        pname, random.choice(CUSTOMERS_POOL),
                        qty, price, qty * price,
                        random.choice(REGIONS),
                        random.choice(["张三", "李四", "王五", "赵六"]),
                    ])
                elif "监控" in topic:
                    writer.writerow([
                        f"DEV-{row_idx:04d}",
                        f"2025-{random.randint(1,12):02d}-{random.randint(1,28):02d} {random.randint(0,23):02d}:{random.randint(0,59):02d}:00",
                        round(random.uniform(15, 45), 1),
                        round(random.uniform(30, 90), 1),
                        round(random.uniform(210, 240), 1),
                        round(random.uniform(0.5, 10), 2),
                        round(random.uniform(100, 2000), 1),
                        random.choice(["正常", "正常", "正常", "警告", "异常"]),
                        random.choice(["无", "低", "中", "高"]),
                    ])
                elif "反馈" in topic:
                    writer.writerow([
                        f"FB-{row_idx:05d}",
                        f"2025-{random.randint(1,12):02d}-{random.randint(1,28):02d}",
                        random.choice(CUSTOMERS_POOL),
                        random.choice([p[0] for p in PRODUCTS_POOL]),
                        random.randint(1, 5),
                        random.choice(["功能建议", "Bug报告", "使用咨询", "投诉", "好评"]),
                        random.choice(["待处理", "处理中", "已解决", "已关闭"]),
                        random.choice(["张三", "李四", "王五"]),
                        random.choice(["已回复", "需跟进", "已升级", ""]),
                    ])
                elif "库存" in topic:
                    stock = random.randint(0, 500)
                    writer.writerow([
                        f"SKU-{row_idx:05d}",
                        random.choice([p[0] for p in PRODUCTS_POOL]),
                        random.choice(["深圳仓", "上海仓", "北京仓", "成都仓"]),
                        stock, random.randint(50, 200),
                        random.randint(0, 100),
                        random.choice(["供应商A", "供应商B", "供应商C"]),
                        f"2025-{random.randint(1,12):02d}-{random.randint(1,28):02d}",
                        "正常" if stock > 50 else "需补货",
                    ])
                elif "项目" in topic:
                    writer.writerow([
                        f"PRJ-{row_idx:04d}",
                        random.choice(["智慧水务", "智能电网", "工业物联", "智慧城市", "物流追踪"]),
                        random.choice(["需求分析", "方案设计", "开发实施", "测试验收", "上线运维"]),
                        f"{random.randint(10, 100)}%",
                        random.choice(["张三", "李四", "王五"]),
                        f"2025-{random.randint(1,6):02d}-01",
                        f"2025-{random.randint(7,12):02d}-30",
                        random.choice(["正常", "正常", "正常", "延期风险", "已延期"]),
                        random.choice(["低", "中", "高"]),
                    ])
                elif "财务" in topic:
                    writer.writerow([
                        f"FIN-{row_idx:06d}",
                        f"2025-{random.randint(1,12):02d}-{random.randint(1,28):02d}",
                        random.choice(["收入", "支出", "转账", "退款"]),
                        random.randint(1000, 500000),
                        random.choice(["营业收入", "研发支出", "管理费用", "销售费用", "税费"]),
                        random.choice(CUSTOMERS_POOL + ["供应商A", "供应商B"]),
                        random.choice(["项目款", "设备采购", "服务费", "工资", "租金"]),
                        random.choice(["张三", "李四"]),
                        random.choice(["已审批", "待审批", "已驳回"]),
                    ])
                elif "物流" in topic:
                    writer.writerow([
                        f"TRK-{row_idx:06d}",
                        f"2025-{random.randint(1,12):02d}-{random.randint(1,28):02d}",
                        random.choice(["北京", "上海", "广州", "深圳", "成都", "武汉"]),
                        random.choice(["电子设备", "配件", "耗材", "文件"]),
                        round(random.uniform(0.5, 500), 1),
                        round(random.uniform(0.01, 5), 2),
                        random.choice(["顺丰", "德邦", "中通", "圆通"]),
                        random.randint(20, 2000),
                        random.choice(["已签收", "运输中", "待发货"]),
                    ])
                elif "质量" in topic:
                    std = round(random.uniform(10, 100), 2)
                    dev = round(random.uniform(-5, 5), 2)
                    writer.writerow([
                        f"QC-{row_idx:05d}",
                        f"2025-{random.randint(1,12):02d}-{random.randint(1,28):02d}",
                        f"BATCH-{random.randint(1000, 9999)}",
                        random.choice(["尺寸", "重量", "电阻", "电压", "强度", "硬度"]),
                        std, round(std + dev, 2), dev,
                        "合格" if abs(dev) < 3 else "不合格",
                        random.choice(["张三", "李四", "王五"]),
                    ])
                elif "培训" in topic:
                    writer.writerow([
                        f"TRN-{row_idx:04d}",
                        f"2025-{random.randint(1,12):02d}-{random.randint(1,28):02d}",
                        random.choice(["Python编程", "项目管理", "数据分析", "安全培训", "领导力"]),
                        random.choice(["张三", "李四", "外聘讲师"]),
                        random.randint(10, 100),
                        f"{random.randint(70, 100)}%",
                        round(random.uniform(3.5, 5.0), 1),
                        random.randint(1000, 50000),
                        random.choice(["", "需补考", "已归档"]),
                    ])
                elif "市场" in topic:
                    budget = random.randint(10000, 200000)
                    actual = int(budget * random.uniform(0.7, 1.3))
                    writer.writerow([
                        f"MKT-{row_idx:04d}",
                        f"2025-{random.randint(1,12):02d}-{random.randint(1,28):02d}",
                        random.choice(["线上推广", "展会", "沙龙", "广告投放", "内容营销"]),
                        random.choice(["新品发布", "品牌宣传", "客户答谢", "技术分享"]),
                        budget, actual,
                        random.randint(20, 5000),
                        f"{round(random.uniform(1, 15), 1)}%",
                        round(actual / budget if budget > 0 else 0, 2),
                    ])

        print(f"[CSV {idx+1:02d}/10] {fname} (50 行)")


# ============================================================
# 主入口
# ============================================================
if __name__ == "__main__":
    print("=" * 60)
    print("开始批量生成测试数据（目标：100+ 个文件）")
    print("=" * 60)

    generate_png_batch()
    generate_txt_batch()
    generate_docx_batch()
    generate_pdf_batch()
    generate_xlsx_batch()
    generate_xls_batch()
    generate_csv_batch()

    print("=" * 60)
    print("所有测试数据生成完毕！")
    files = sorted(OUTPUT_DIR.iterdir())
    total = len(files)
    total_size = sum(f.stat().st_size for f in files)
    print(f"输出目录: {OUTPUT_DIR}")
    print(f"文件总数: {total}")
    print(f"总大小: {total_size / 1024 / 1024:.1f} MB")
    print()

    # 按类型统计
    from collections import Counter
    ext_counts = Counter(f.suffix for f in files)
    for ext, count in sorted(ext_counts.items()):
        print(f"  {ext}: {count} 个")
    print("=" * 60)
